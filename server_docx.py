"""MCP Server — Leitura e manipulacao de DOCX (comentarios, citacoes, referencias)."""

import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt
from lxml import etree
from docx.opc.part import Part
from docx.opc.packuri import PackURI
from mcp.server.fastmcp import FastMCP
import datetime

mcp = FastMCP("docx-manager")

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
M_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"
MATHLIKE_RE = re.compile(r"[=+\-*/≤≥≠∀ΣΦϕΘρΔεηαβγμ∣\[\]{}()]|max|min|VaR|CVaR|MAE|MAPE|RMSE|MASE")
EQUATION_END_RE = re.compile(r"\((\d+)\)\s*$")
EQUATION_LABEL_RE = re.compile(r"^\s*\((\d+)\)\s*$")
SINGLE_EQ_REF_RE = re.compile(r"\b(?:Equação|eq\.)\s*\((\d+)\)", re.IGNORECASE)
RANGE_EQ_REF_RE = re.compile(r"\b(?:Equações|eqs\.)\s*\((\d+)\)\s*(?:a|até)\s*\((\d+)\)", re.IGNORECASE)
PAIR_EQ_REF_RE = re.compile(r"\b(?:Equações|eqs\.)\s*\((\d+)\)\s*e\s*\((\d+)\)", re.IGNORECASE)
TABLE_CAPTION_RE = re.compile(r'^Tabela\s+\d+\s+[–-]\s+.+')
TABLE_NUM_RE = re.compile(r'^(Tabela\s+)(\d+)(\s+[–-]\s+.+)')
SINGLE_TABLE_REF_RE = re.compile(r'\bTabela\s+(\d+)\b', re.IGNORECASE)

W14_NS = "http://schemas.microsoft.com/office/word/2010/wordml"
W15_NS = "http://schemas.microsoft.com/office/word/2012/wordml"
W16CID_NS = "http://schemas.microsoft.com/office/word/2016/wordml/cid"
W16CEX_NS = "http://schemas.microsoft.com/office/word/2018/wordml/cex"
CR_NS = "http://schemas.microsoft.com/office/comments/2020/reactions"
RT_COMMENTS_EXTENDED = "http://schemas.microsoft.com/office/2011/relationships/commentsExtended"
RT_COMMENTS_IDS = "http://schemas.microsoft.com/office/2016/09/relationships/commentsIds"
RT_COMMENTS_EXTENSIBLE = "http://schemas.microsoft.com/office/2018/08/relationships/commentsExtensible"
CT_COMMENTS_EXTENDED = "application/vnd.openxmlformats-officedocument.wordprocessingml.commentsExtended+xml"


# ── Helpers ──────────────────────────────────────────────────────────────────


def _get_comments(doc_path: str) -> list[dict]:
    """Extrai comentarios do DOCX via python-docx API."""
    doc = Document(doc_path)
    return [
        {"id": str(c.comment_id), "author": c.author, "text": c.text.strip()}
        for c in doc.comments
    ]


def _get_comment_para_ids(doc_path: str) -> dict:
    """Retorna {paraId: {"done": bool, "paraIdParent": str|None}} de commentsExtended.xml."""
    doc = Document(doc_path)
    try:
        ext_part = doc.part.part_related_by(RT_COMMENTS_EXTENDED)
    except KeyError:
        return {}
    tree = etree.fromstring(ext_part.blob)
    result = {}
    for ex in tree.findall(f"{{{W15_NS}}}commentEx"):
        para_id = ex.get(f"{{{W15_NS}}}paraId", "")
        done = ex.get(f"{{{W15_NS}}}done", "0") == "1"
        parent = ex.get(f"{{{W15_NS}}}paraIdParent")
        result[para_id] = {"done": done, "paraIdParent": parent}
    return result


def _get_comment_situations(doc_path: str) -> dict:
    """Retorna {comment_id: resolved} mapeando ID visivel do comentario ao status resolvido."""
    doc = Document(doc_path)
    doc_part = doc.part

    # Build paraId -> done from commentsExtended.xml
    try:
        ext_part = doc_part.part_related_by(RT_COMMENTS_EXTENDED)
    except KeyError:
        return {}
    tree = etree.fromstring(ext_part.blob)
    para_done = {}
    for ex in tree.findall(f"{{{W15_NS}}}commentEx"):
        pid = ex.get(f"{{{W15_NS}}}paraId", "")
        para_done[pid] = ex.get(f"{{{W15_NS}}}done", "0") == "1"

    # Build paraId -> comment_id from comments.xml
    comments_elm = doc_part._comments_part.element
    para_to_cid = {}
    for c in comments_elm.findall(qn("w:comment")):
        para = c.find(f".//{{{W_NS}}}p")
        pid = para.get(f"{{{W14_NS}}}paraId") if para is not None else ""
        if pid:
            para_to_cid[pid] = c.get(qn("w:id"))

    return {cid: para_done.get(pid, False) for pid, cid in para_to_cid.items()}


def _get_thumbsup_para_ids(doc_part) -> set:
    """Retorna paraIds de comentarios com reacao 👍 (reactionType='1') em commentsExtensible.xml."""
    try:
        ids_part = doc_part.part_related_by(RT_COMMENTS_IDS)
    except KeyError:
        return set()
    ids_tree = etree.fromstring(ids_part.blob)
    para_to_durable = {
        el.get(f"{{{W16CID_NS}}}paraId"): el.get(f"{{{W16CID_NS}}}durableId")
        for el in ids_tree
        if el.get(f"{{{W16CID_NS}}}paraId") and el.get(f"{{{W16CID_NS}}}durableId")
    }
    durable_to_para = {v: k for k, v in para_to_durable.items()}

    try:
        ext_part = doc_part.part_related_by(RT_COMMENTS_EXTENSIBLE)
    except KeyError:
        return set()
    ext_tree = etree.fromstring(ext_part.blob)
    thumbsup = set()
    for entry in ext_tree.findall(f"{{{W16CEX_NS}}}commentExtensible"):
        durable_id = entry.get(f"{{{W16CEX_NS}}}durableId")
        reactions = entry.findall(f".//{{{CR_NS}}}reaction")
        if any(r.get("reactionType") == "1" for r in reactions):
            para_id = durable_to_para.get(durable_id)
            if para_id:
                thumbsup.add(para_id)
    return thumbsup


def _get_paragraph_text(para) -> str:
    return "".join(run.text for run in para.runs if run.text)


def _get_paragraph_full_text(para) -> str:
    parts = []
    for node in para._element.iter():
        if isinstance(node.tag, str) and node.tag.endswith("}t") and node.text:
            parts.append(node.text)
    return "".join(parts)


def _set_run_font(run, font_name: str = "Times New Roman", font_size: int = 10, bold: bool | None = None) -> None:
    run.font.name = font_name
    run.font.size = Pt(font_size)
    if bold is not None:
        run.font.bold = bold

    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    for key in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        r_fonts.set(qn(key), font_name)


def _set_paragraph_font(paragraph, font_name: str = "Times New Roman", font_size: int = 12) -> None:
    """Aplica fonte padrao a todos os runs do paragrafo.

    Se o paragrafo estiver vazio, cria um run vazio apenas para persistir a formatacao.
    """
    if not paragraph.runs:
        paragraph.add_run("")
    for run in paragraph.runs:
        _set_run_font(run, font_name=font_name, font_size=font_size, bold=run.font.bold)


def _word_text_nodes(para):
    return [node for node in para._element.iter() if getattr(node, "tag", None) == qn("w:t")]


def _replace_in_word_text_nodes(para, transform) -> bool:
    nodes = _word_text_nodes(para)
    if not nodes:
        return False

    original = "".join(node.text or "" for node in nodes)
    updated = transform(original)
    if updated == original:
        return False

    nodes[0].text = updated
    for node in nodes[1:]:
        node.text = ""
    return True


def _paragraph_has_omml(para) -> bool:
    for node in para._element.iter():
        if getattr(node, "tag", None) in {f"{{{M_NS}}}oMath", f"{{{M_NS}}}oMathPara"}:
            return True
    return False


def _paragraph_is_in_table(para) -> bool:
    parent = para._element.getparent()
    while parent is not None:
        if parent.tag == qn("w:tc"):
            return True
        parent = parent.getparent()
    return False


def _get_body_children(doc: Document):
    return list(doc._body._element)


def _paragraph_text_from_element(element) -> str:
    parts = []
    for node in element.iter():
        if getattr(node, "tag", None) == qn("w:t") and node.text:
            parts.append(node.text)
    return "".join(parts).strip()


def _get_table_caption(doc: Document, table) -> str:
    children = _get_body_children(doc)
    table_el = table._element
    child_idx = children.index(table_el)

    for idx in range(child_idx - 1, -1, -1):
        element = children[idx]
        if element.tag == qn("w:p"):
            text = _paragraph_text_from_element(element)
            if text:
                return text
    return ""


def _get_table_caption_paragraph(doc: Document, table):
    children = _get_body_children(doc)
    table_el = table._element
    child_idx = children.index(table_el)

    for idx in range(child_idx - 1, -1, -1):
        element = children[idx]
        if element.tag == qn("w:p"):
            text = _paragraph_text_from_element(element)
            if text:
                for paragraph in doc.paragraphs:
                    if paragraph._element is element:
                        return paragraph
                return None
    return None


def _set_table_caption(doc: Document, table, caption_text: str) -> None:
    existing = _get_table_caption_paragraph(doc, table)
    if existing is not None and TABLE_CAPTION_RE.match(_get_paragraph_full_text(existing).strip()):
        if existing.runs:
            existing.runs[0].text = caption_text
            for run in existing.runs[1:]:
                run.text = ""
        else:
            existing.add_run(caption_text)
        for run in existing.runs:
            _set_run_font(run, font_name="Times New Roman", font_size=10, bold=False)
        return

    p = OxmlElement("w:p")
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = caption_text
    r.append(t)
    p.append(r)
    table._element.addprevious(p)


def _cell_border_sides(cell_el) -> set[str]:
    tc_pr = cell_el.find(qn("w:tcPr"))
    if tc_pr is None:
        return set()
    tc_borders = tc_pr.find(qn("w:tcBorders"))
    if tc_borders is None:
        return set()

    sides = set()
    for side in ("top", "bottom", "left", "right"):
        border = tc_borders.find(qn(f"w:{side}"))
        if border is not None and border.get(qn("w:val")) not in (None, "nil"):
            sides.add(side)
    return sides


def _set_cell_border(cell, **kwargs) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)

    for edge in ("top", "bottom", "left", "right"):
        data = kwargs.get(edge)
        if data is None:
            continue
        tag = qn(f"w:{edge}")
        element = tc_borders.find(tag)
        if element is None:
            element = OxmlElement(f"w:{edge}")
            tc_borders.append(element)
        for key, value in data.items():
            element.set(qn(f"w:{key}"), str(value))


def _clear_top_bottom_borders(table) -> None:
    for row in table.rows:
        for cell in row.cells:
            _set_cell_border(
                cell,
                top={"val": "nil"},
                bottom={"val": "nil"},
            )


def _apply_minimal_table_borders(table) -> None:
    if not table.rows:
        return

    _clear_top_bottom_borders(table)

    for cell in table.rows[0].cells:
        _set_cell_border(
            cell,
            top={"val": "single", "sz": "4", "color": "000000"},
            bottom={"val": "single", "sz": "4", "color": "000000"},
        )

    for cell in table.rows[-1].cells:
        _set_cell_border(
            cell,
            bottom={"val": "single", "sz": "4", "color": "000000"},
        )


def _apply_table_font(table, font_name: str = "Times New Roman", font_size: int = 10) -> None:
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    _set_run_font(run, font_name=font_name, font_size=font_size, bold=(row_idx == 0))


def _parse_table_numbers(table_numbers: str, total_tables: int) -> list[int]:
    if not table_numbers.strip():
        return list(range(1, total_tables + 1))

    result = []
    for part in table_numbers.split(","):
        value = int(part.strip())
        if value < 1 or value > total_tables:
            raise ValueError(f"Tabela {value} fora do intervalo 1..{total_tables}")
        result.append(value)
    return sorted(set(result))


def _remove_table_row(table, row_index: int) -> None:
    row = table.rows[row_index]
    row._tr.getparent().remove(row._tr)


def _ensure_paragraph_anchor(paragraph):
    if paragraph.runs:
        return paragraph.runs[0]
    return paragraph.add_run("")


def _add_comment_to_table(doc: Document, table, text: str, author: str = "docx-manager") -> bool:
    if not table.rows or not table.rows[0].cells:
        return False
    paragraph = table.rows[0].cells[0].paragraphs[0]
    anchor = _ensure_paragraph_anchor(paragraph)
    doc.add_comment(anchor, text, author=author)
    return True


def _table_report(doc: Document) -> list[dict]:
    report = []
    for idx, table in enumerate(doc.tables, start=1):
        caption = _get_table_caption(doc, table)
        jc = None
        layout = None
        tbl_pr = table._element.find(qn("w:tblPr"))
        if tbl_pr is not None:
            jc_el = tbl_pr.find(qn("w:jc"))
            if jc_el is not None:
                jc = jc_el.get(qn("w:val"))
            layout_el = tbl_pr.find(qn("w:tblLayout"))
            if layout_el is not None:
                layout = layout_el.get(qn("w:type"))

        rows = table._element.findall(qn("w:tr"))
        first_row_cells = rows[0].findall(qn("w:tc")) if rows else []
        last_row_cells = rows[-1].findall(qn("w:tc")) if rows else []
        middle_rows = rows[1:-1] if len(rows) > 2 else []

        issues = []
        if not TABLE_CAPTION_RE.match(caption):
            issues.append('legenda acima ausente ou fora do padrão "Tabela N – ..."')
        if jc != "center":
            issues.append("tabela sem alinhamento explícito centralizado")
        if layout == "fixed":
            issues.append("layout fixo; não está em autoajuste")

        first_ok = bool(first_row_cells) and all(
            {"top", "bottom"}.issubset(_cell_border_sides(cell))
            for cell in first_row_cells
        )
        last_ok = bool(last_row_cells) and all(
            "bottom" in _cell_border_sides(cell)
            for cell in last_row_cells
        )
        if not first_ok:
            issues.append("primeira linha sem bordas superior e inferior em todas as células")
        if not last_ok:
            issues.append("última linha sem borda inferior em todas as células")

        has_internal_borders = False
        for row in middle_rows:
            for cell in row.findall(qn("w:tc")):
                sides = _cell_border_sides(cell)
                if "bottom" in sides:
                    has_internal_borders = True
                    break
            if has_internal_borders:
                break
        if has_internal_borders:
            issues.append("há bordas internas além do padrão mínimo esperado")

        font_issues = []
        for row_idx, row in enumerate(table.rows):
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        if not run.text:
                            continue
                        if run.font.name not in (None, "Times New Roman"):
                            font_issues.append("texto fora de Times New Roman")
                            break
                        if run.font.size is not None and abs(run.font.size.pt - 10) > 0.1:
                            font_issues.append("texto fora do tamanho 10")
                            break
                        if row_idx == 0 and run.font.bold is False:
                            font_issues.append("primeira linha sem negrito completo")
                            break
                    if font_issues:
                        break
                if font_issues:
                    break
            if font_issues:
                break
        issues.extend(font_issues)

        report.append({
            "table_index": idx,
            "status": "ok" if not issues else "erro",
            "issues": issues,
            "summary": "conforme esperado" if not issues else "; ".join(issues),
            "caption": caption,
        })
    return report


def _looks_like_equation_text(text: str) -> bool:
    if not text:
        return False
    return bool(MATHLIKE_RE.search(text))


def _extract_equations(doc: Document) -> list[dict]:
    equations = []
    for idx, para in enumerate(doc.paragraphs):
        if _paragraph_is_in_table(para):
            continue
        full_text = _get_paragraph_full_text(para).strip()
        plain_text = _get_paragraph_text(para).strip()
        match = EQUATION_END_RE.search(full_text)
        if not match:
            continue
        if not _paragraph_has_omml(para):
            continue

        eq_number = int(match.group(1))
        eq_text = EQUATION_END_RE.sub("", full_text).strip()
        if not eq_text or not EQUATION_LABEL_RE.fullmatch(plain_text):
            continue

        equations.append({
            "equation_number": eq_number,
            "paragraph_index": idx,
            "equation_text": eq_text,
            "label_text": plain_text,
        })
    return equations


def _extract_equation_refs(text: str) -> list[dict]:
    refs = []

    for match in RANGE_EQ_REF_RE.finditer(text):
        start = int(match.group(1))
        end = int(match.group(2))
        refs.append({
            "kind": "range",
            "raw": match.group(0),
            "numbers": list(range(start, end + 1)) if start <= end else [],
            "start": start,
            "end": end,
        })

    for match in PAIR_EQ_REF_RE.finditer(text):
        refs.append({
            "kind": "pair",
            "raw": match.group(0),
            "numbers": [int(match.group(1)), int(match.group(2))],
        })

    for match in SINGLE_EQ_REF_RE.finditer(text):
        refs.append({
            "kind": "single",
            "raw": match.group(0),
            "numbers": [int(match.group(1))],
        })

    return refs


def _validate_equation_state(doc: Document) -> dict:
    equations = _extract_equations(doc)
    numbers = [item["equation_number"] for item in equations]
    unique_numbers = sorted(set(numbers))
    duplicates = sorted(n for n in unique_numbers if numbers.count(n) > 1)

    expected_sequence = []
    missing_numbers = []
    if unique_numbers:
        expected_sequence = list(range(min(unique_numbers), max(unique_numbers) + 1))
        missing_numbers = [n for n in expected_sequence if n not in unique_numbers]

    reference_items = []
    invalid_references = []
    invalid_ranges = []

    for idx, para in enumerate(doc.paragraphs):
        text = _get_paragraph_full_text(para)
        for ref in _extract_equation_refs(text):
            entry = {
                "paragraph_index": idx,
                "raw": ref["raw"],
                "kind": ref["kind"],
                "numbers": ref["numbers"],
            }
            reference_items.append(entry)

            if ref["kind"] == "range" and ref["start"] > ref["end"]:
                invalid_ranges.append(entry | {"start": ref["start"], "end": ref["end"]})

            for n in ref["numbers"]:
                if n not in unique_numbers:
                    invalid_references.append(entry | {"missing_equation_number": n})

    return {
        "equation_count": len(equations),
        "equations": equations,
        "numbers_in_order": numbers,
        "duplicates": duplicates,
        "missing_numbers_in_sequence": missing_numbers,
        "references": reference_items,
        "invalid_references": invalid_references,
        "invalid_ranges": invalid_ranges,
        "is_valid": not duplicates and not invalid_references and not invalid_ranges,
    }


def _renumber_reference_text(
    text: str,
    mapping_options: dict[int, list[int]],
    fallback_old_number: int | None = None,
) -> tuple[str, list[dict]]:
    warnings: list[dict] = []

    def resolve_number(value: int) -> int | None:
        options = mapping_options.get(value, [])
        if len(options) == 1:
            return options[0]
        # Corrige uma referencia orfa quando o texto aponta para a equacao
        # exibida imediatamente acima, mas o numero ficou defasado em 1.
        if not options and fallback_old_number is not None and abs(value - fallback_old_number) == 1:
            fallback_options = mapping_options.get(fallback_old_number, [])
            if len(fallback_options) == 1:
                return fallback_options[0]
        return None

    def register_warning(kind: str, raw: str, numbers: list[int], reason: str) -> None:
        warnings.append({
            "kind": kind,
            "raw": raw,
            "numbers": numbers,
            "reason": reason,
        })

    def replace_range(match):
        start = int(match.group(1))
        end = int(match.group(2))
        resolved_start = resolve_number(start)
        resolved_end = resolve_number(end)
        if resolved_start is None or resolved_end is None:
            register_warning("range", match.group(0), [start, end], "intervalo com mapeamento ambiguo ou ausente")
            return match.group(0)
        connector = "até" if "até" in match.group(0).lower() else "a"
        prefix = "eqs." if match.group(0).lower().startswith("eqs.") else "Equações"
        return f"{prefix} ({resolved_start}) {connector} ({resolved_end})"

    def replace_pair(match):
        left = int(match.group(1))
        right = int(match.group(2))
        resolved_left = resolve_number(left)
        resolved_right = resolve_number(right)
        if resolved_left is None or resolved_right is None:
            register_warning("pair", match.group(0), [left, right], "par com mapeamento ambiguo ou ausente")
            return match.group(0)
        prefix = "eqs." if match.group(0).lower().startswith("eqs.") else "Equações"
        return f"{prefix} ({resolved_left}) e ({resolved_right})"

    def replace_single(match):
        value = int(match.group(1))
        resolved = resolve_number(value)
        if resolved is None:
            register_warning("single", match.group(0), [value], "referencia com mapeamento ambiguo ou ausente")
            return match.group(0)
        prefix = "eq." if match.group(0).lower().startswith("eq.") else "Equação"
        return f"{prefix} ({resolved})"

    updated = RANGE_EQ_REF_RE.sub(replace_range, text)
    updated = PAIR_EQ_REF_RE.sub(replace_pair, updated)
    updated = SINGLE_EQ_REF_RE.sub(replace_single, updated)
    return updated, warnings


# ── Tools ────────────────────────────────────────────────────────────────────


@mcp.tool()
def list_paragraphs(doc_path: str) -> list[dict]:
    """Lista paragrafos do DOCX com indice e texto."""
    doc = Document(doc_path)
    return [
        {"index": i, "text": _get_paragraph_text(p)[:200]}
        for i, p in enumerate(doc.paragraphs)
        if _get_paragraph_text(p).strip()
    ]


@mcp.tool()
def list_comments(doc_path: str) -> list[dict]:
    """Lista todos os comentarios do DOCX."""
    return _get_comments(doc_path)


@mcp.tool()
def get_situation(doc_path: str, comment_id: str = "") -> dict:
    """Retorna se um comentario esta resolvido ou nao. Se comment_id for omitido, retorna todos."""
    situations = _get_comment_situations(doc_path)
    if comment_id:
        if comment_id not in situations:
            return {"error": f"Comentario {comment_id} nao encontrado."}
        return {"id": comment_id, "resolved": situations[comment_id]}
    return [{"id": cid, "resolved": resolved} for cid, resolved in situations.items()]


@mcp.tool()
def list_equations(doc_path: str) -> list[dict]:
    """Lista as equacoes encontradas no DOCX com numero, paragrafo e texto linearizado."""
    doc = Document(doc_path)
    return _extract_equations(doc)


@mcp.tool()
def validate_equation_references(doc_path: str) -> dict:
    """Valida a numeracao das equacoes e as referencias textuais do documento."""
    doc = Document(doc_path)
    return _validate_equation_state(doc)


@mcp.tool()
def validate_tables_format(doc_path: str) -> list[dict]:
    """Valida legenda, alinhamento, bordas e fonte das tabelas do DOCX."""
    doc = Document(doc_path)
    return _table_report(doc)


@mcp.tool()
def report_tables_format(doc_path: str) -> list[str]:
    """Retorna um relatório textual de ok/erro por tabela."""
    doc = Document(doc_path)
    report = _table_report(doc)
    return [
        f"Tabela {item['table_index']}: {item['status']} - {item['summary']}"
        for item in report
    ]


@mcp.tool()
def find_citar_paragraphs(doc_path: str) -> list[dict]:
    """Encontra paragrafos que possuem comentario 'citar'. Retorna indice + texto do paragrafo."""
    doc = Document(doc_path)
    comments = _get_comments(doc_path)
    citar_ids = {c["id"] for c in comments if c["text"].strip().lower() == "citar"}

    if not citar_ids:
        return [{"info": "Nenhum comentario 'citar' encontrado."}]

    # Mapear commentRangeStart no XML do documento
    body = doc.element.body
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    results = []

    for para_idx, para in enumerate(doc.paragraphs):
        para_xml = para._element
        starts = para_xml.findall(".//w:commentRangeStart", ns)
        for s in starts:
            cid = s.get(f'{{{ns["w"]}}}id')
            if cid in citar_ids:
                results.append({
                    "paragraph_index": para_idx,
                    "text": _get_paragraph_text(para)[:500],
                    "comment_id": cid,
                })

    return results if results else [{"info": "Nenhum paragrafo marcado com 'citar'."}]


@mcp.tool()
def find_instruction_paragraphs(doc_path: str) -> list[dict]:
    """Encontra paragrafos com comentarios de instrucao (qualquer texto exceto 'citar')."""
    doc = Document(doc_path)
    comments = _get_comments(doc_path)
    instruction_ids = {
        c["id"]: c["text"]
        for c in comments
        if c["text"].strip().lower() != "citar"
    }

    if not instruction_ids:
        return [{"info": "Nenhum comentario de instrucao encontrado."}]

    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    results = []

    for para_idx, para in enumerate(doc.paragraphs):
        starts = para._element.findall(".//w:commentRangeStart", ns)
        for s in starts:
            cid = s.get(f'{{{ns["w"]}}}id')
            if cid in instruction_ids:
                results.append({
                    "paragraph_index": para_idx,
                    "text": _get_paragraph_text(para)[:500],
                    "comment_id": cid,
                    "instruction": instruction_ids[cid],
                })

    return results if results else [{"info": "Nenhum paragrafo com instrucao encontrado."}]


@mcp.tool()
def reply_comment(
    doc_path: str,
    comment_id: str,
    reply_text: str,
    output_path: str = "",
) -> str:
    """Responde a um comentario existente com uma mensagem em thread OOXML."""
    import secrets

    save_path = output_path if output_path and output_path != doc_path else doc_path

    w14 = "http://schemas.microsoft.com/office/word/2010/wordml"
    w15 = "http://schemas.microsoft.com/office/word/2012/wordml"
    RT_EXT = "http://schemas.microsoft.com/office/2011/relationships/commentsExtended"
    CT_EXT = "application/vnd.openxmlformats-officedocument.wordprocessingml.commentsExtended+xml"

    doc = Document(doc_path)
    doc_part = doc.part

    # Access comments.xml via CommentsPart (lxml element, auto-serialized on save)
    comments_elm = doc_part._comments_part.element

    parent_comment = None
    max_id = 0
    parent_para_id = None
    for c in comments_elm.findall(qn("w:comment")):
        cid = c.get(qn("w:id"))
        max_id = max(max_id, int(cid))
        if cid == comment_id:
            parent_comment = c
            parent_para_id = c.get(qn("w14:paraId")) or c.get(f"{{{W_NS}}}paraId") or c.get("paraId", "")

    if parent_comment is None:
        return f"Comentario {comment_id} nao encontrado."

    new_id = str(max_id + 1)
    new_para_id = secrets.token_hex(4).upper()
    now = datetime.datetime.now(datetime.timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ")

    # Create new comment element with lxml
    new_comment = etree.SubElement(comments_elm, qn("w:comment"))
    new_comment.set(qn("w:id"), new_id)
    new_comment.set(qn("w:author"), "docx-manager")
    new_comment.set(qn("w:date"), now)
    new_comment.set(f"{{{w14}}}paraId", new_para_id)
    new_p = etree.SubElement(new_comment, qn("w:p"))
    new_p.set(f"{{{w14}}}paraId", secrets.token_hex(4).upper())
    new_r = etree.SubElement(new_p, qn("w:r"))
    new_t = etree.SubElement(new_r, qn("w:t"))
    new_t.text = reply_text

    # Access/create commentsExtended.xml via OPC
    try:
        ext_part = doc_part.part_related_by(RT_EXT)
        ext_tree = etree.fromstring(ext_part.blob)
    except KeyError:
        ext_part = None
        ext_tree = etree.Element(f"{{{w15}}}commentsEx")

    new_ex = etree.SubElement(ext_tree, f"{{{w15}}}commentEx")
    new_ex.set(f"{{{w15}}}paraId", new_para_id)
    new_ex.set(f"{{{w15}}}done", "0")
    if parent_para_id:
        new_ex.set(f"{{{w15}}}paraIdParent", parent_para_id)

    ext_blob = etree.tostring(ext_tree, xml_declaration=True, encoding="UTF-8", standalone=True)
    if ext_part is None:
        from docx.opc.part import Part
        from docx.opc.packuri import PackURI
        ext_part = Part(
            PackURI("/word/commentsExtended.xml"),
            CT_EXT,
            ext_blob,
            doc_part.package,
        )
        doc_part.relate_to(ext_part, RT_EXT)
    else:
        ext_part._blob = ext_blob

    # Add comment range markers to document.xml
    doc_elm = doc_part.element
    body = doc_elm.find(qn("w:body"))
    inserted = False
    for p in body.iter(qn("w:p")):
        for el in p:
            if el.tag == qn("w:commentRangeStart") and el.get(qn("w:id")) == comment_id:
                start_el = etree.SubElement(p, qn("w:commentRangeStart"))
                start_el.set(qn("w:id"), new_id)
                end_el = etree.SubElement(p, qn("w:commentRangeEnd"))
                end_el.set(qn("w:id"), new_id)
                ref_run = etree.SubElement(p, qn("w:r"))
                ref_el = etree.SubElement(ref_run, qn("w:commentReference"))
                ref_el.set(qn("w:id"), new_id)
                inserted = True
                break
        if inserted:
            break

    # Save (ZIP_DEFLATED automatic, Content_Types and rels managed by python-docx)
    if output_path and output_path != doc_path:
        Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(save_path)

    return f"Resposta inserida no comentario {comment_id} -> {save_path}"


@mcp.tool()
def remove_resolved_comments(doc_path: str, output_path: str = "") -> str:
    """Remove comentarios resolvidos com logica de cascata em threads."""
    save_path = output_path if output_path and output_path != doc_path else doc_path

    w15 = "http://schemas.microsoft.com/office/word/2012/wordml"
    RT_EXT = "http://schemas.microsoft.com/office/2011/relationships/commentsExtended"

    doc = Document(doc_path)
    doc_part = doc.part

    # Access commentsExtended.xml via OPC
    try:
        ext_part = doc_part.part_related_by(RT_EXT)
    except KeyError:
        return "0 comentarios resolvidos removidos (commentsExtended.xml ausente)."

    ext_tree = etree.fromstring(ext_part.blob)

    # Build para_info: {paraId: {done, parent}}
    para_info = {}
    for ex in ext_tree.findall(f"{{{w15}}}commentEx"):
        pid = ex.get(f"{{{w15}}}paraId", "")
        para_info[pid] = {
            "done": ex.get(f"{{{w15}}}done", "0") == "1",
            "parent": ex.get(f"{{{w15}}}paraIdParent"),
        }

    # Tratar comentarios com 👍 como resolvidos
    thumbsup_ids = _get_thumbsup_para_ids(doc_part)
    for pid in thumbsup_ids:
        if pid in para_info:
            para_info[pid]["done"] = True

    if not any(info["done"] for info in para_info.values()):
        return "0 comentarios resolvidos removidos."

    # Map paraId -> comment_id from comments.xml
    # paraId is on the child <w:p> paragraph inside the comment, not on <w:comment> itself
    comments_elm = doc_part._comments_part.element
    W14_NS = "http://schemas.microsoft.com/office/word/2010/wordml"
    para_to_cid = {}
    for c in comments_elm.findall(qn("w:comment")):
        para = c.find(f".//{{{W_NS}}}p")
        pid = para.get(f"{{{W14_NS}}}paraId") if para is not None else ""
        if pid:
            para_to_cid[pid] = c.get(qn("w:id"))

    # Group threads: parent_paraId -> [child_paraIds]
    threads = {}
    children_set = set()
    for pid, info in para_info.items():
        if info["parent"]:
            threads.setdefault(info["parent"], []).append(pid)
            children_set.add(pid)

    # Determine which comment IDs to remove (cascade logic)
    ids_to_remove = set()

    for parent_pid, child_pids in threads.items():
        parent_info = para_info.get(parent_pid, {})
        if parent_info.get("done"):
            # Rule 1: parent resolved -> remove entire thread
            if parent_pid in para_to_cid:
                ids_to_remove.add(para_to_cid[parent_pid])
            for cpid in child_pids:
                if cpid in para_to_cid:
                    ids_to_remove.add(para_to_cid[cpid])
        else:
            resolved_children = [cp for cp in child_pids if para_info.get(cp, {}).get("done")]
            unresolved_children = [cp for cp in child_pids if not para_info.get(cp, {}).get("done")]
            for cpid in resolved_children:
                if cpid in para_to_cid:
                    ids_to_remove.add(para_to_cid[cpid])
            # Rule 3: no unresolved children left -> also remove parent
            if not unresolved_children and parent_pid in para_to_cid:
                ids_to_remove.add(para_to_cid[parent_pid])

    # Standalone resolved (not in any thread)
    for pid, info in para_info.items():
        if info["done"] and pid not in children_set and pid not in threads:
            if pid in para_to_cid:
                ids_to_remove.add(para_to_cid[pid])

    if not ids_to_remove:
        return "0 comentarios resolvidos removidos."

    # Remove from comments.xml
    for c in list(comments_elm.findall(qn("w:comment"))):
        if c.get(qn("w:id")) in ids_to_remove:
            comments_elm.remove(c)

    # Remove from document.xml (lxml getparent — fixes O(n²) issue)
    doc_elm = doc_part.element
    for tag in ("commentRangeStart", "commentRangeEnd"):
        for el in list(doc_elm.iter(qn(f"w:{tag}"))):
            if el.get(qn("w:id")) in ids_to_remove:
                el.getparent().remove(el)

    # Remove commentReference (also remove orphan w:r)
    for ref in list(doc_elm.iter(qn("w:commentReference"))):
        if ref.get(qn("w:id")) in ids_to_remove:
            run = ref.getparent()
            run.remove(ref)
            if len(run) == 0 and run.text is None:
                run.getparent().remove(run)

    # Remove from commentsExtended.xml
    cid_to_para = {v: k for k, v in para_to_cid.items()}
    removed_para_ids = {cid_to_para[cid] for cid in ids_to_remove if cid in cid_to_para}
    for ex in list(ext_tree.findall(f"{{{w15}}}commentEx")):
        if ex.get(f"{{{w15}}}paraId", "") in removed_para_ids:
            ext_tree.remove(ex)
    ext_part._blob = etree.tostring(ext_tree, xml_declaration=True, encoding="UTF-8", standalone=True)

    # Save with ZIP_DEFLATED (automatic via python-docx)
    if output_path and output_path != doc_path:
        Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(save_path)

    return f"{len(ids_to_remove)} comentarios resolvidos removidos -> {save_path}"


@mcp.tool()
def insert_candidate_comment(
    doc_path: str,
    paragraph_index: int,
    candidate_text: str,
    output_path: str = "",
    author: str = "docx-manager",
) -> str:
    """Insere um comentario de candidato em um paragrafo especifico do DOCX.

    Args:
        doc_path: caminho do DOCX de entrada
        paragraph_index: indice do paragrafo
        candidate_text: texto do comentario (ex: 'CANDIDATO: Olivares et al. (2024)\\nScore: 0.91')
        output_path: caminho do DOCX de saida (default: salva in-place)
        author: autor do comentario (ex: 'claude', 'codex', 'docx-manager')
    """
    doc = Document(doc_path)
    if paragraph_index >= len(doc.paragraphs):
        return f"Indice {paragraph_index} fora do range (total: {len(doc.paragraphs)})"

    para = doc.paragraphs[paragraph_index]
    anchor = _ensure_paragraph_anchor(para)
    doc.add_comment(anchor, candidate_text, author=author)

    save_path = output_path if output_path and output_path != doc_path else doc_path
    if output_path and output_path != doc_path:
        Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(save_path)
    return f"Comentario inserido no paragrafo {paragraph_index} -> {save_path}"


@mcp.tool()
def insert_citation(
    doc_path: str,
    paragraph_index: int,
    citation: str,
    reference: str,
    source: str = "",
    output_path: str = "",
    author: str = "docx-manager",
) -> str:
    """Insere citacao inline no paragrafo e adiciona referencia ao final do documento.

    Args:
        doc_path: caminho do DOCX
        paragraph_index: indice do paragrafo
        citation: texto da citacao inline, ex: '(OLIVARES et al., 2024)'
        reference: referencia completa para o final do documento
        source: DOI, link ou path do documento citado (opcional, incluido no comentario)
        output_path: caminho de saida (default: salva in-place)
        author: autor do comentario (ex: 'claude', 'codex', 'docx-manager')
    """
    doc = Document(doc_path)
    if paragraph_index >= len(doc.paragraphs):
        return f"Indice {paragraph_index} fora do range"

    para = doc.paragraphs[paragraph_index]

    # Adicionar citacao ao final do paragrafo e normalizar fonte do paragrafo inteiro
    para.add_run(f" {citation}")
    _set_paragraph_font(para, font_name="Times New Roman", font_size=12)

    # Comentario unificado com citacao e fonte
    if source:
        anchor = _ensure_paragraph_anchor(para)
        comment_text = f"Citação: {citation}\nFonte: {source}"
        doc.add_comment(anchor, comment_text, author=author)

    # Adicionar referencia ao final do documento (sem duplicar)
    existing_refs = {_get_paragraph_text(p).strip() for p in doc.paragraphs}
    if reference.strip() not in existing_refs:
        ref_para = doc.add_paragraph(reference)
        _set_paragraph_font(ref_para, font_name="Times New Roman", font_size=12)

    save_path = output_path if output_path and output_path != doc_path else doc_path
    if output_path and output_path != doc_path:
        Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(save_path)
    return f"Citacao '{citation}' inserida no paragrafo {paragraph_index}. Referencia adicionada ao final."


@mcp.tool()
def replace_paragraph_text(
    doc_path: str,
    paragraph_index: int,
    new_text: str,
    comment: str,
    output_path: str = "",
) -> str:
    """Substitui o texto de um paragrafo e insere comentario de atribuicao.

    Args:
        doc_path: caminho do DOCX
        paragraph_index: indice do paragrafo
        new_text: novo texto para o paragrafo
        comment: texto do comentario de atribuicao
        output_path: caminho de saida (default: salva in-place)
    """
    doc = Document(doc_path)
    if paragraph_index >= len(doc.paragraphs):
        return f"Indice {paragraph_index} fora do range (total: {len(doc.paragraphs)})"

    para = doc.paragraphs[paragraph_index]

    # Capturar apenas o estado de negrito do primeiro run antes de limpar.
    # Para texto substituido, a regra e padronizar em Times New Roman 12.
    font_name = "Times New Roman"
    font_size = 12
    bold = None
    if para.runs:
        first_run = para.runs[0]
        bold = first_run.font.bold

    # Limpar runs existentes e inserir novo texto
    if para.runs:
        para.runs[0].text = new_text
        for run in para.runs[1:]:
            run.text = ""
        target_run = para.runs[0]
    else:
        target_run = para.add_run(new_text)

    _set_run_font(target_run, font_name=font_name, font_size=font_size, bold=bold)
    _set_paragraph_font(para, font_name=font_name, font_size=font_size)

    # Comentario de atribuicao
    anchor = _ensure_paragraph_anchor(para)
    doc.add_comment(anchor, comment, author="docx-manager")

    save_path = output_path if output_path and output_path != doc_path else doc_path
    if output_path and output_path != doc_path:
        Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(save_path)
    return f"Paragrafo {paragraph_index} substituido e comentado -> {save_path}"


@mcp.tool()
def apply_table_style(
    doc_path: str,
    table_numbers: str = "",
    font_name: str = "Times New Roman",
    font_size: int = 10,
    caption_prefix: str = "",
    output_path: str = "",
) -> dict:
    """Aplica estilo padrao de tabelas e valida o resultado.

    `table_numbers` deve ser uma string CSV com indices 1-based, ex.: "3,4,8".
    Se vazio, aplica em todas as tabelas. `output_path` default: salva in-place.
    """
    doc = Document(doc_path)
    selected = _parse_table_numbers(table_numbers, len(doc.tables))
    before = _table_report(doc)

    for table_number in selected:
        table = doc.tables[table_number - 1]
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = True
        _apply_minimal_table_borders(table)
        _apply_table_font(table, font_name=font_name, font_size=font_size)
        if caption_prefix:
            _set_table_caption(doc, table, f"Tabela {table_number} – {caption_prefix} {table_number}".strip())

    save_path = output_path if output_path and output_path != doc_path else doc_path
    if output_path and output_path != doc_path:
        Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(save_path)

    after_doc = Document(save_path)
    after = _table_report(after_doc)

    return {
        "ok": all(item["status"] == "ok" for item in after if item["table_index"] in selected),
        "message": "Estilo de tabela aplicado." if selected else "Nenhuma tabela selecionada.",
        "table_numbers": selected,
        "validation_before": [item for item in before if item["table_index"] in selected],
        "validation_after": [item for item in after if item["table_index"] in selected],
        "output_path": save_path,
    }


@mcp.tool()
def set_table_caption(
    doc_path: str,
    table_number: int,
    caption_text: str,
    output_path: str = "",
) -> dict:
    """Insere ou corrige a legenda acima de uma tabela específica. `output_path` default: salva in-place."""
    doc = Document(doc_path)
    if table_number < 1 or table_number > len(doc.tables):
        return {"ok": False, "message": f"Tabela {table_number} fora do intervalo 1..{len(doc.tables)}"}

    table = doc.tables[table_number - 1]
    _set_table_caption(doc, table, caption_text)
    save_path = output_path if output_path and output_path != doc_path else doc_path
    if output_path and output_path != doc_path:
        Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(save_path)
    after_doc = Document(save_path)
    report = _table_report(after_doc)
    item = next(x for x in report if x["table_index"] == table_number)
    return {
        "ok": "legenda acima ausente ou fora do padrão \"Tabela N – ...\"" not in item["issues"],
        "table_number": table_number,
        "caption": caption_text,
        "validation_after": item,
        "output_path": save_path,
    }


@mcp.tool()
def validate_tables_with_comments(
    doc_path: str,
    output_path: str = "",
) -> dict:
    """Valida as tabelas, insere comentario nas que tiverem erro e retorna o relatorio textual. `output_path` default: salva in-place."""
    doc = Document(doc_path)
    report = _table_report(doc)
    annotated = []

    for item in report:
        if item["status"] != "erro":
            continue
        table = doc.tables[item["table_index"] - 1]
        comment_text = f"ERRO DE FORMATAÇÃO DA TABELA {item['table_index']}: {item['summary']}"
        if _add_comment_to_table(doc, table, comment_text):
            annotated.append(item["table_index"])

    save_path = output_path if output_path and output_path != doc_path else doc_path
    if output_path and output_path != doc_path:
        Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(save_path)

    return {
        "output_path": save_path,
        "annotated_tables": annotated,
        "report": [
            f"Tabela {item['table_index']}: {item['status']} - {item['summary']}"
            for item in report
        ],
    }


@mcp.tool()
def renumber_equations(
    doc_path: str,
    start_number: int = 1,
    output_path: str = "",
) -> dict:
    """Renumera equacoes em ordem de aparicao e atualiza referencias textuais.

    Considera apenas equacoes OMML em paragrafo isolado com rotulo `(N)`.
    A operacao renumera em ordem de aparicao e tenta atualizar referencias
    textuais de forma best effort. Quando houver ambiguidade ou falta de
    mapeamento, retorna avisos para inspecao manual. `output_path` default:
    salva in-place.
    """
    doc = Document(doc_path)
    before = _validate_equation_state(doc)

    if not before["equations"]:
        return {
            "ok": False,
            "message": "Nenhuma equacao encontrada no documento.",
            "validation_before": before,
        }

    mapping: dict[int, int] = {}
    mapping_options: dict[int, list[int]] = {}
    equation_by_paragraph = {
        eq["paragraph_index"]: eq["equation_number"]
        for eq in before["equations"]
    }
    duplicate_reference_numbers = set(before["duplicates"])

    for offset, eq in enumerate(before["equations"]):
        old_number = eq["equation_number"]
        new_number = start_number + offset
        mapping[old_number] = new_number
        mapping_options.setdefault(old_number, []).append(new_number)

    # Atualizar labels das equacoes sem tocar nos elementos matematicos.
    for offset, eq in enumerate(before["equations"]):
        para = doc.paragraphs[eq["paragraph_index"]]
        old_number = eq["equation_number"]
        new_number = start_number + offset

        def relabel(text: str) -> str:
            return re.sub(rf"\({old_number}\)\s*$", f"({new_number})", text)

        _replace_in_word_text_nodes(para, relabel)

    # Atualizar referencias textuais no documento inteiro.
    warnings: list[dict] = []
    for idx, para in enumerate(doc.paragraphs):
        fallback_old_number = None
        prev_idx = idx - 1
        while prev_idx >= 0:
            prev_text = _get_paragraph_full_text(doc.paragraphs[prev_idx]).strip()
            if prev_text:
                fallback_old_number = equation_by_paragraph.get(prev_idx)
                break
            prev_idx -= 1
        nodes = _word_text_nodes(para)
        if not nodes:
            continue
        original = "".join(node.text or "" for node in nodes)
        updated, local_warnings = _renumber_reference_text(
            original,
            mapping_options,
            fallback_old_number=fallback_old_number,
        )
        if updated != original:
            nodes[0].text = updated
            for node in nodes[1:]:
                node.text = ""
        for item in local_warnings:
            item["paragraph_index"] = idx
            item["paragraph_text"] = original[:300]
            warnings.append(item)

    for number in sorted(duplicate_reference_numbers):
        warnings.append({
            "paragraph_index": None,
            "kind": "duplicate-equation-number",
            "raw": f"({number})",
            "numbers": [number],
            "reason": "numero de equacao repetido no documento; referencias para esse numero podem exigir inspecao manual",
        })

    for item in before["invalid_ranges"]:
        warnings.append({
            "paragraph_index": item["paragraph_index"],
            "kind": "invalid-range",
            "raw": item["raw"],
            "numbers": item["numbers"],
            "reason": "intervalo textual invalido; inspecao manual recomendada",
        })

    save_path = output_path if output_path and output_path != doc_path else doc_path
    if output_path and output_path != doc_path:
        Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(save_path)

    after_doc = Document(save_path)
    after = _validate_equation_state(after_doc)

    return {
        "ok": True,
        "message": "Equacoes renumeradas. Referencias textuais foram atualizadas quando o mapeamento era inequivoco.",
        "mapping": mapping,
        "warnings": warnings,
        "validation_before": before,
        "validation_after": after,
        "output_path": save_path,
    }


@mcp.tool()
def renumber_tables(
    doc_path: str,
    start_number: int = 1,
    output_path: str = "",
) -> dict:
    """Renumera tabelas em ordem de aparicao e atualiza referencias textuais.

    A renumeracao normaliza todas as referencias para 'Tabela' com inicial maiuscula.
    Valida o estado antes e depois. `output_path` default: salva in-place.
    """
    doc = Document(doc_path)
    before = _table_report(doc)

    if not doc.tables:
        return {
            "ok": False,
            "message": "Nenhuma tabela encontrada no documento.",
            "validation_before": before,
        }

    # Construir mapping: número atual -> novo número
    mapping = {}
    caption_paras = []
    for idx, table in enumerate(doc.tables):
        cap_para = _get_table_caption_paragraph(doc, table)
        if cap_para is None:
            continue
        cap_text = _get_paragraph_full_text(cap_para).strip()
        match = TABLE_NUM_RE.match(cap_text)
        if not match:
            continue
        old_num = int(match.group(2))
        new_num = start_number + len(mapping)
        mapping[old_num] = new_num
        caption_paras.append((cap_para, old_num, new_num))

    if not mapping:
        return {
            "ok": False,
            "message": "Nenhuma legenda no formato 'Tabela N – ...' encontrada.",
            "validation_before": before,
        }

    # Atualizar legendas
    for cap_para, old_num, new_num in caption_paras:
        def relabel(text, _old=old_num, _new=new_num):
            return TABLE_NUM_RE.sub(lambda m: f"{m.group(1)}{_new}{m.group(3)}", text)
        _replace_in_word_text_nodes(cap_para, relabel)

    # Atualizar referencias textuais no documento inteiro
    def replace_table_refs(text):
        def sub_ref(m):
            n = int(m.group(1))
            if n in mapping:
                return f"Tabela {mapping[n]}"
            return m.group(0)
        return SINGLE_TABLE_REF_RE.sub(sub_ref, text)

    for para in doc.paragraphs:
        # Evitar re-processar legendas (já atualizadas)
        if any(para is cp for cp, _, _ in caption_paras):
            continue
        _replace_in_word_text_nodes(para, replace_table_refs)

    save_path = output_path if output_path and output_path != doc_path else doc_path
    if output_path and output_path != doc_path:
        Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(save_path)

    after_doc = Document(save_path)
    after = _table_report(after_doc)

    return {
        "ok": all(item["status"] == "ok" for item in after),
        "message": "Tabelas renumeradas e referencias atualizadas.",
        "mapping": mapping,
        "validation_before": before,
        "validation_after": after,
        "output_path": save_path,
    }


if __name__ == "__main__":
    mcp.run()
