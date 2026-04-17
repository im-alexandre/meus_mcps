"""Test script for new server_docx.py features."""
import base64
import datetime
import os
import sys
import traceback
import tempfile
import shutil
from pathlib import Path
from zipfile import ZipFile, ZIP_DEFLATED
import xml.etree.ElementTree as ET

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree

# Ensure server_docx is importable
sys.path.insert(0, os.path.dirname(__file__))

import server_docx as sd

RESULTS = []
TMPDIR = tempfile.mkdtemp(prefix="test_docx_")
TINY_PNG_BASE64 = "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAusB9sXv0WQAAAAASUVORK5CYII="


def report(name, passed, detail=""):
    status = "PASS" if passed else "FAIL"
    RESULTS.append((name, passed))
    msg = f"  [{status}] {name}"
    if detail:
        msg += f" — {detail}"
    print(msg)


def make_simple_docx(path):
    """Create a basic DOCX with paragraphs, a table, and an equation."""
    doc = Document()
    doc.add_paragraph("Introduction paragraph.")
    doc.add_paragraph("This is the body text. See Tabela 1 for details.")
    doc.add_paragraph("Tabela 1 \u2013 Sample data")
    table = doc.add_table(rows=3, cols=2)
    table.cell(0, 0).text = "Header A"
    table.cell(0, 1).text = "Header B"
    table.cell(1, 0).text = "1"
    table.cell(1, 1).text = "2"
    table.cell(2, 0).text = "3"
    table.cell(2, 1).text = "4"
    doc.add_paragraph("")
    doc.add_paragraph("Tabela 2 \u2013 More data")
    table2 = doc.add_table(rows=2, cols=2)
    table2.cell(0, 0).text = "X"
    table2.cell(0, 1).text = "Y"
    table2.cell(1, 0).text = "10"
    table2.cell(1, 1).text = "20"
    doc.add_paragraph("See Tabela 2 for more.")
    doc.add_paragraph("Conclusion paragraph.")
    doc.save(path)


def make_docx_with_comment(path):
    """Create a DOCX with a comment (requires OOXML manipulation)."""
    doc = Document()
    p1 = doc.add_paragraph("Paragraph with instruction comment.")
    p2 = doc.add_paragraph("Paragraph marked for citation.")
    p3 = doc.add_paragraph("Regular paragraph.")

    # Add comments using python-docx add_comment
    anchor1 = p1.runs[0] if p1.runs else p1.add_run("")
    doc.add_comment(anchor1, "Reescrever em tom formal", author="reviewer")

    anchor2 = p2.runs[0] if p2.runs else p2.add_run("")
    doc.add_comment(anchor2, "citar", author="reviewer")

    doc.save(path)


def make_docx_with_inline_figure(path):
    """Create a DOCX with one inline figure anchored to a paragraph."""
    image_path = os.path.join(TMPDIR, f"{Path(path).stem}_figure.png")
    with open(image_path, "wb") as handle:
        handle.write(base64.b64decode(TINY_PNG_BASE64))

    doc = Document()
    doc.add_paragraph("Paragraph before figure.")
    figure_paragraph = doc.add_paragraph()
    figure_paragraph.add_run().add_picture(image_path)
    doc.add_paragraph("Paragraph after figure.")
    doc.save(path)


def inject_comment_state(path, done=False, thumbsup=False, comment_index=0):
    """Inject commentsExtended/commentsIds/commentsExtensible metadata into a DOCX."""
    from io import BytesIO

    w_ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    w14_ns = "http://schemas.microsoft.com/office/word/2010/wordml"
    w15_ns = "http://schemas.microsoft.com/office/word/2012/wordml"
    w16cid_ns = "http://schemas.microsoft.com/office/word/2016/wordml/cid"
    w16cex_ns = "http://schemas.microsoft.com/office/word/2018/wordml/cex"
    cr_ns = "http://schemas.microsoft.com/office/comments/2020/reactions"
    ct_ns = "http://schemas.openxmlformats.org/package/2006/content-types"
    rel_ns = "http://schemas.openxmlformats.org/package/2006/relationships"

    with open(path, "rb") as handle:
        data = BytesIO(handle.read())

    with ZipFile(data, "r") as zin:
        parts = {name: zin.read(name) for name in zin.namelist()}

    ET.register_namespace("w", w_ns)
    ET.register_namespace("w14", w14_ns)
    ET.register_namespace("w15", w15_ns)
    ET.register_namespace("w16cid", w16cid_ns)
    ET.register_namespace("w16cex", w16cex_ns)
    ET.register_namespace("cr", cr_ns)
    ET.register_namespace("", rel_ns)

    comments_tree = ET.fromstring(parts["word/comments.xml"])
    comments = comments_tree.findall(f"{{{w_ns}}}comment")
    comment = comments[comment_index]
    para_id = comment.get(f"{{{w14_ns}}}paraId") or f"AABB{comment_index + 1:04d}"
    comment.set(f"{{{w14_ns}}}paraId", para_id)
    comment_paragraph = comment.find(f"{{{w_ns}}}p")
    if comment_paragraph is not None:
        comment_paragraph.set(f"{{{w14_ns}}}paraId", para_id)
    parts["word/comments.xml"] = ET.tostring(comments_tree, xml_declaration=True, encoding="UTF-8")

    def ensure_override(ct_tree, part_name, content_type):
        for override in ct_tree.findall(f"{{{ct_ns}}}Override"):
            if override.get("PartName") == part_name:
                override.set("ContentType", content_type)
                return
        override = ET.SubElement(ct_tree, f"{{{ct_ns}}}Override")
        override.set("PartName", part_name)
        override.set("ContentType", content_type)

    def ensure_relationship(rels_tree, rel_type, target):
        for relationship in rels_tree.findall(f"{{{rel_ns}}}Relationship"):
            if relationship.get("Type") == rel_type and relationship.get("Target") == target:
                return
        existing_ids = [rel.get("Id", "") for rel in rels_tree.findall(f"{{{rel_ns}}}Relationship")]
        next_id = max(
            (
                int(rel_id.replace("rId", ""))
                for rel_id in existing_ids
                if rel_id.startswith("rId") and rel_id[3:].isdigit()
            ),
            default=0,
        ) + 1
        relationship = ET.SubElement(rels_tree, f"{{{rel_ns}}}Relationship")
        relationship.set("Id", f"rId{next_id}")
        relationship.set("Type", rel_type)
        relationship.set("Target", target)

    ct_tree = ET.fromstring(parts["[Content_Types].xml"])
    rels_path = "word/_rels/document.xml.rels"
    rels_tree = ET.fromstring(parts[rels_path])

    if done:
        ext_root = ET.Element(f"{{{w15_ns}}}commentsEx")
        ex = ET.SubElement(ext_root, f"{{{w15_ns}}}commentEx")
        ex.set(f"{{{w15_ns}}}paraId", para_id)
        ex.set(f"{{{w15_ns}}}done", "1")
        parts["word/commentsExtended.xml"] = ET.tostring(ext_root, xml_declaration=True, encoding="UTF-8")
        ensure_override(
            ct_tree,
            "/word/commentsExtended.xml",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.commentsExtended+xml",
        )
        ensure_relationship(rels_tree, sd.RT_COMMENTS_EXTENDED, "commentsExtended.xml")

    if thumbsup:
        durable_id = f"DURABLE{comment_index + 1:04d}"
        ids_root = ET.Element(f"{{{w16cid_ns}}}commentsIds")
        comment_id = ET.SubElement(ids_root, f"{{{w16cid_ns}}}commentId")
        comment_id.set(f"{{{w16cid_ns}}}paraId", para_id)
        comment_id.set(f"{{{w16cid_ns}}}durableId", durable_id)
        parts["word/commentsIds.xml"] = ET.tostring(ids_root, xml_declaration=True, encoding="UTF-8")

        ext_root = ET.Element(f"{{{w16cex_ns}}}commentsExtensible")
        extensible = ET.SubElement(ext_root, f"{{{w16cex_ns}}}commentExtensible")
        extensible.set(f"{{{w16cex_ns}}}durableId", durable_id)
        reactions = ET.SubElement(extensible, f"{{{cr_ns}}}reactions")
        reaction = ET.SubElement(reactions, f"{{{cr_ns}}}reaction")
        reaction.set("reactionType", "1")
        parts["word/commentsExtensible.xml"] = ET.tostring(ext_root, xml_declaration=True, encoding="UTF-8")

        ensure_override(
            ct_tree,
            "/word/commentsIds.xml",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.commentsIds+xml",
        )
        ensure_override(
            ct_tree,
            "/word/commentsExtensible.xml",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.commentsExtensible+xml",
        )
        ensure_relationship(rels_tree, sd.RT_COMMENTS_IDS, "commentsIds.xml")
        ensure_relationship(rels_tree, sd.RT_COMMENTS_EXTENSIBLE, "commentsExtensible.xml")

    parts["[Content_Types].xml"] = ET.tostring(ct_tree, xml_declaration=True, encoding="UTF-8")
    parts[rels_path] = ET.tostring(rels_tree, xml_declaration=True, encoding="UTF-8")

    buf = BytesIO()
    with ZipFile(buf, "w", ZIP_DEFLATED) as zout:
        for name, content in parts.items():
            zout.writestr(name, content)
    with open(path, "wb") as handle:
        handle.write(buf.getvalue())


def make_docx_with_equations(path):
    """Create a DOCX with equation-like paragraphs."""
    doc = Document()
    doc.add_paragraph("Introduction. See Equa\u00e7\u00e3o (1) and Equa\u00e7\u00e3o (2).")
    p1 = doc.add_paragraph()
    omath1 = etree.fromstring(
        """
        <m:oMath xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">
          <m:r><m:t>x = a + b</m:t></m:r>
        </m:oMath>
        """
    )
    p1._element.append(omath1)
    p1.add_run(" (1)")
    doc.add_paragraph("Some text between equations.")
    p2 = doc.add_paragraph()
    omath2 = etree.fromstring(
        """
        <m:oMath xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">
          <m:r><m:t>y = c * d</m:t></m:r>
        </m:oMath>
        """
    )
    p2._element.append(omath2)
    p2.add_run(" (2)")
    doc.add_paragraph("Conclusion referencing eq. (1).")
    doc.save(path)


def make_docx_with_orphan_equation_reference(path):
    """Create a DOCX where a paragraph after an equation points to an adjacent missing number."""
    doc = Document()
    doc.add_paragraph("Texto introdut\u00f3rio.")
    p1 = doc.add_paragraph()
    omath1 = etree.fromstring(
        """
        <m:oMath xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">
          <m:r><m:t>x = a + b</m:t></m:r>
        </m:oMath>
        """
    )
    p1._element.append(omath1)
    p1.add_run(" (1)")
    doc.add_paragraph("Pela Equa\u00e7\u00e3o (2), obt\u00e9m-se o resultado final.")
    doc.save(path)


def make_docx_with_duplicate_equation_numbers(path):
    """Create a DOCX with duplicate equation labels to ensure best-effort renumbering."""
    doc = Document()
    for label in (1, 1):
        p = doc.add_paragraph()
        omath = etree.fromstring(
            """
            <m:oMath xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">
              <m:r><m:t>x = a + b</m:t></m:r>
            </m:oMath>
            """
        )
        p._element.append(omath)
        p.add_run(f" ({label})")
    doc.add_paragraph("Ver Equação (1).")
    doc.save(path)


# ────────────────────────────────────────────────────────
# Test 1: list_comments
# ────────────────────────────────────────────────────────
def test_list_comments_context():
    path = os.path.join(TMPDIR, "test_comments.docx")
    make_docx_with_comment(path)
    try:
        result = sd.list_comments(path)
        report("list_comments: returns comments", len(result) == 2, str(result))
        first = next((item for item in result if item.get("text", "").startswith("Reescrever")), {})
        second = next((item for item in result if item.get("text", "").lower() == "citar"), {})
        report("list_comments: first paragraph index", first.get("paragraph_index") == 0, str(first))
        report(
            "list_comments: first paragraph text",
            first.get("paragraph_text", "").startswith("Paragraph with instruction comment."),
            str(first),
        )
        report("list_comments: second paragraph index", second.get("paragraph_index") == 1, str(second))
    except Exception as e:
        report("list_comments context", False, str(e))
        traceback.print_exc()


def test_removed_public_apis():
    report("removed api: get_situation", not hasattr(sd, "get_situation"))
    report("removed api: find_citar_paragraphs", not hasattr(sd, "find_citar_paragraphs"))
    report("removed api: find_instruction_paragraphs", not hasattr(sd, "find_instruction_paragraphs"))


def test_list_comments_resolved_by_thumbsup():
    path = os.path.join(TMPDIR, "test_comments_thumbsup.docx")
    make_docx_with_comment(path)
    try:
        inject_comment_state(path, thumbsup=True)
        result = sd.list_comments(path)
        first = next((item for item in result if item.get("paragraph_index") == 0), {})
        report("list_comments: thumbsup resolved", first.get("resolved") is True, str(first))
    except Exception as e:
        report("list_comments thumbsup", False, str(e))
        traceback.print_exc()


# ────────────────────────────────────────────────────────
# Test 2: reply_comment
# ────────────────────────────────────────────────────────
def test_reply_comment():
    path = os.path.join(TMPDIR, "test_reply.docx")
    out_path = os.path.join(TMPDIR, "test_reply_out.docx")
    make_docx_with_comment(path)
    try:
        comments = sd._get_comments(path)
        if not comments:
            report("reply_comment", False, "No comments in test doc")
            return
        cid = comments[0]["id"]
        result = sd.reply_comment(path, cid, "Done, rewritten.", out_path)
        report("reply_comment: returns success", "Resposta inserida" in result, result)

        # Verify the reply exists in comments.xml
        comments_after = sd._get_comments(out_path)
        reply_found = any(c["text"] == "Done, rewritten." for c in comments_after)
        report("reply_comment: reply in comments.xml", reply_found)
        reply = next((c for c in comments_after if c["text"] == "Done, rewritten."), None)
        report("reply_comment: reply linked to parent", bool(reply and reply.get("parent_id") == cid), str(reply))

        # Verify commentsExtended.xml has the new entry
        ext = sd._get_comment_para_ids(out_path)
        report("reply_comment: commentsExtended updated", len(ext) > 0)

        tree = sd.list_comments(out_path)
        parent = next((c for c in tree if c["id"] == cid), None)
        nested_reply = next((r for r in (parent or {}).get("replies", []) if r["text"] == "Done, rewritten."), None)
        report("list_comments: returns roots only", all(not c.get("parent_id") for c in tree), str(tree))
        report("list_comments: nests replies", nested_reply is not None)
        report("list_comments: reply inherits paragraph index", (nested_reply or {}).get("paragraph_index") == 0, str(nested_reply))
    except Exception as e:
        report("reply_comment", False, str(e))
        traceback.print_exc()


# ────────────────────────────────────────────────────────
# Test 3: replace_paragraph_text
# ────────────────────────────────────────────────────────
def test_replace_paragraph_text():
    path = os.path.join(TMPDIR, "test_replace.docx")
    out_path = os.path.join(TMPDIR, "test_replace_out.docx")
    doc = Document()
    doc.add_paragraph("Original text here.")
    doc.add_paragraph("Another paragraph.")
    doc.save(path)
    try:
        result = sd.replace_paragraph_text(path, 0, "Replaced text.", "AI rewrite", out_path)
        report("replace_paragraph_text: returns success", "substituido" in result.lower(), result)

        doc2 = Document(out_path)
        new_text = sd._get_paragraph_text(doc2.paragraphs[0])
        report("replace_paragraph_text: text changed", new_text.strip() == "Replaced text.")
        run0 = doc2.paragraphs[0].runs[0]
        report(
            "replace_paragraph_text: font normalized",
            run0.font.name == "Times New Roman" and round(run0.font.size.pt, 1) == 12.0,
            f"name={run0.font.name}, size={None if run0.font.size is None else run0.font.size.pt}",
        )

        # Check comment was added
        comments = sd._get_comments(out_path)
        has_attr = any("AI rewrite" in c["text"] for c in comments)
        report("replace_paragraph_text: attribution comment", has_attr)
    except Exception as e:
        report("replace_paragraph_text", False, str(e))
        traceback.print_exc()


# ────────────────────────────────────────────────────────
# Test 4: remove_resolved_comments
# ────────────────────────────────────────────────────────
def test_remove_resolved_comments():
    path = os.path.join(TMPDIR, "test_resolved.docx")
    out_path = os.path.join(TMPDIR, "test_resolved_out.docx")
    make_docx_with_comment(path)
    try:
        # First check it has comments
        before = sd._get_comments(path)
        report("remove_resolved_comments: has comments before", len(before) > 0)

        # No commentsExtended.xml in our test doc, so nothing to remove
        result = sd.remove_resolved_comments(path, out_path)
        report("remove_resolved_comments: handles missing extended", "0 comentarios" in result or "ausente" in result, result)
    except Exception as e:
        report("remove_resolved_comments", False, str(e))
        traceback.print_exc()


def test_remove_resolved_comments_with_resolved():
    """Test with a manually-crafted DOCX that has commentsExtended.xml with done=1."""
    path = os.path.join(TMPDIR, "test_resolved2.docx")
    out_path = os.path.join(TMPDIR, "test_resolved2_out.docx")
    make_docx_with_comment(path)

    try:
        inject_comment_state(path, done=True)

        comments_before = sd._get_comments(path)
        count_before = len(comments_before)

        result = sd.remove_resolved_comments(path, out_path)
        report("remove_resolved_comments (resolved): removes comment", "1 comentarios" in result, result)

        comments_after = sd._get_comments(out_path)
        report("remove_resolved_comments (resolved): count decreased", len(comments_after) < count_before,
               f"before={count_before}, after={len(comments_after)}")

    except Exception as e:
        report("remove_resolved_comments (resolved)", False, str(e))
        traceback.print_exc()


def test_remove_resolved_comments_with_thumbsup():
    path = os.path.join(TMPDIR, "test_resolved_thumbsup.docx")
    out_path = os.path.join(TMPDIR, "test_resolved_thumbsup_out.docx")
    make_docx_with_comment(path)

    try:
        inject_comment_state(path, thumbsup=True)

        comments_before = sd._get_comments(path)
        count_before = len(comments_before)

        result = sd.remove_resolved_comments(path, out_path)
        report("remove_resolved_comments (thumbsup): removes comment", "1 comentarios" in result, result)

        comments_after = sd._get_comments(out_path)
        report(
            "remove_resolved_comments (thumbsup): count decreased",
            len(comments_after) < count_before,
            f"before={count_before}, after={len(comments_after)}",
        )
    except Exception as e:
        report("remove_resolved_comments (thumbsup)", False, str(e))
        traceback.print_exc()


# ────────────────────────────────────────────────────────
# Test 5: renumber_tables
# ────────────────────────────────────────────────────────
def test_renumber_tables():
    path = os.path.join(TMPDIR, "test_renum_tables.docx")
    out_path = os.path.join(TMPDIR, "test_renum_tables_out.docx")
    make_simple_docx(path)
    try:
        result = sd.renumber_tables(path, start_number=1, output_path=out_path)
        report("renumber_tables: returns ok", result.get("ok", False) or result.get("mapping"), str(result.get("message", "")))
        report("renumber_tables: has mapping", bool(result.get("mapping")))

        if result.get("mapping"):
            doc2 = Document(out_path)
            # Check captions were preserved
            found_captions = []
            for p in doc2.paragraphs:
                text = sd._get_paragraph_full_text(p).strip()
                if text.startswith("Tabela"):
                    found_captions.append(text)
            report("renumber_tables: captions present", len(found_captions) >= 2, str(found_captions))
    except Exception as e:
        report("renumber_tables", False, str(e))
        traceback.print_exc()


# ────────────────────────────────────────────────────────
# Test 6: insert_citation with source param
# ────────────────────────────────────────────────────────
def test_insert_citation():
    path = os.path.join(TMPDIR, "test_citation.docx")
    out_path = os.path.join(TMPDIR, "test_citation_out.docx")
    doc = Document()
    doc.add_paragraph("This finding is important.")
    doc.add_paragraph("Another paragraph.")
    doc.save(path)
    try:
        result = sd.insert_citation(
            path, 0,
            "(SILVA, 2020)",
            "SILVA, J. Artigo completo. Journal X, 2020.",
            source="https://doi.org/10.1234/test",
            output_path=out_path,
        )
        report("insert_citation: returns success", "Citacao" in result, result)

        doc2 = Document(out_path)
        p_text = sd._get_paragraph_text(doc2.paragraphs[0])
        report("insert_citation: citation in text", "(SILVA, 2020)" in p_text, p_text)
        para_runs = [r for r in doc2.paragraphs[0].runs if r.text.strip()]
        para_font_ok = all(r.font.name == "Times New Roman" and round(r.font.size.pt, 1) == 12.0 for r in para_runs)
        report("insert_citation: paragraph font normalized", para_font_ok)

        # Check reference added at end
        last_text = sd._get_paragraph_text(doc2.paragraphs[-1])
        report("insert_citation: reference at end", "SILVA" in last_text, last_text)
        last_runs = [r for r in doc2.paragraphs[-1].runs if r.text.strip()]
        ref_font_ok = all(r.font.name == "Times New Roman" and round(r.font.size.pt, 1) == 12.0 for r in last_runs)
        report("insert_citation: reference font normalized", ref_font_ok)

        # Check source comment
        comments = sd._get_comments(out_path)
        has_source = any("doi.org" in c["text"] for c in comments)
        report("insert_citation: source comment", has_source)

        # Test deduplication: insert same citation again
        result2 = sd.insert_citation(
            out_path, 1,
            "(SILVA, 2020)",
            "SILVA, J. Artigo completo. Journal X, 2020.",
            output_path=out_path,
        )
        doc3 = Document(out_path)
        ref_count = sum(1 for p in doc3.paragraphs
                        if "SILVA, J. Artigo completo" in sd._get_paragraph_text(p))
        report("insert_citation: no duplicate reference", ref_count == 1, f"ref_count={ref_count}")
    except Exception as e:
        report("insert_citation", False, str(e))
        traceback.print_exc()


# ────────────────────────────────────────────────────────
# Test 7: output_path optional (in-place save)
# ────────────────────────────────────────────────────────
def test_output_path_optional():
    path = os.path.join(TMPDIR, "test_inplace.docx")
    doc = Document()
    doc.add_paragraph("Test paragraph for in-place save.")
    doc.save(path)
    try:
        result = sd.replace_paragraph_text(path, 0, "Updated.", "auto", output_path="")
        report("output_path optional: in-place works", os.path.exists(path))
        doc2 = Document(path)
        text = sd._get_paragraph_text(doc2.paragraphs[0])
        report("output_path optional: text updated in-place", "Updated." in text, text)
    except Exception as e:
        report("output_path optional", False, str(e))
        traceback.print_exc()


# ────────────────────────────────────────────────────────
# Test 8: renumber_equations
# ────────────────────────────────────────────────────────
def test_renumber_equations():
    path = os.path.join(TMPDIR, "test_renum_eq.docx")
    out_path = os.path.join(TMPDIR, "test_renum_eq_out.docx")
    make_docx_with_equations(path)
    try:
        result = sd.renumber_equations(path, start_number=1, output_path=out_path)
        report("renumber_equations: ok or valid", result.get("ok", False) or "mapping" in result, str(result.get("message", "")))
    except Exception as e:
        report("renumber_equations", False, str(e))
        traceback.print_exc()


def test_renumber_equations_with_orphan_reference():
    path = os.path.join(TMPDIR, "test_renum_eq_orphan.docx")
    out_path = os.path.join(TMPDIR, "test_renum_eq_orphan_out.docx")
    make_docx_with_orphan_equation_reference(path)
    try:
        result = sd.renumber_equations(path, start_number=1, output_path=out_path)
        report("renumber_equations orphan: ok", result.get("ok", False), str(result.get("message", "")))
        after = result.get("validation_after", {})
        report("renumber_equations orphan: validation ok", after.get("is_valid", False))
        doc = Document(out_path)
        text = sd._get_paragraph_full_text(doc.paragraphs[2])
        report("renumber_equations orphan: reference corrected", "Equação (1)" in text, text)
    except Exception as e:
        report("renumber_equations orphan", False, str(e))
        traceback.print_exc()


def test_renumber_equations_with_duplicate_numbers():
    path = os.path.join(TMPDIR, "test_renum_eq_dupe.docx")
    out_path = os.path.join(TMPDIR, "test_renum_eq_dupe_out.docx")
    make_docx_with_duplicate_equation_numbers(path)
    try:
        result = sd.renumber_equations(path, start_number=1, output_path=out_path)
        report("renumber_equations duplicate: ok", result.get("ok", False), str(result.get("message", "")))
        warnings = result.get("warnings", [])
        has_duplicate_warning = any(item.get("kind") == "duplicate-equation-number" for item in warnings)
        report("renumber_equations duplicate: warns", has_duplicate_warning, str(warnings))
        after = result.get("validation_after", {})
        report("renumber_equations duplicate: renumbered sequentially", after.get("numbers_in_order") == [1, 2], str(after.get("numbers_in_order")))
    except Exception as e:
        report("renumber_equations duplicate", False, str(e))
        traceback.print_exc()


# ────────────────────────────────────────────────────────
# Test 9: apply_table_style
# ────────────────────────────────────────────────────────
def test_apply_table_style():
    path = os.path.join(TMPDIR, "test_style.docx")
    out_path = os.path.join(TMPDIR, "test_style_out.docx")
    make_simple_docx(path)
    try:
        result = sd.apply_table_style(path, table_numbers="1", output_path=out_path)
        report("apply_table_style: returns result", isinstance(result, dict))
        report("apply_table_style: has validation", "validation_after" in result)
    except Exception as e:
        report("apply_table_style", False, str(e))
        traceback.print_exc()


# ────────────────────────────────────────────────────────
# Test 10: set_table_caption
# ────────────────────────────────────────────────────────
def test_set_table_caption():
    path = os.path.join(TMPDIR, "test_cap.docx")
    out_path = os.path.join(TMPDIR, "test_cap_out.docx")
    make_simple_docx(path)
    try:
        result = sd.set_table_caption(path, 1, "Updated caption", out_path)
        report("set_table_caption: ok", result.get("ok", False), str(result))
        doc2 = Document(out_path)
        caption = sd._get_table_caption(doc2, doc2.tables[0])
        report("set_table_caption: normalizes prefix", caption == "Tabela 1 – Updated caption", caption)
    except Exception as e:
        report("set_table_caption", False, str(e))
        traceback.print_exc()


# ────────────────────────────────────────────────────────
# Test 11: set_table_source
# ────────────────────────────────────────────────────────
def test_set_table_source():
    path = os.path.join(TMPDIR, "test_table_source.docx")
    out_path = os.path.join(TMPDIR, "test_table_source_out.docx")
    make_simple_docx(path)
    try:
        result = sd.set_table_source(path, 1, output_path=out_path)
        expected = f"Fonte: Autor ({datetime.datetime.now().year})"
        report("set_table_source: ok", result.get("ok", False), str(result))
        doc2 = Document(out_path)
        texts = [sd._get_paragraph_full_text(p).strip() for p in doc2.paragraphs if sd._get_paragraph_full_text(p).strip()]
        report("set_table_source: default source inserted", expected in texts, str(texts))
    except Exception as e:
        report("set_table_source", False, str(e))
        traceback.print_exc()


# ────────────────────────────────────────────────────────
# Test 12: set_figure_caption
# ────────────────────────────────────────────────────────
def test_set_figure_caption():
    path = os.path.join(TMPDIR, "test_figure_caption.docx")
    out_path = os.path.join(TMPDIR, "test_figure_caption_out.docx")
    make_docx_with_inline_figure(path)
    try:
        result = sd.set_figure_caption(path, 1, "Fluxo do modelo", out_path)
        report("set_figure_caption: ok", result.get("ok", False), str(result))
        doc2 = Document(out_path)
        texts = [sd._get_paragraph_full_text(p).strip() for p in doc2.paragraphs if sd._get_paragraph_full_text(p).strip()]
        report("set_figure_caption: normalizes prefix", "Figura 1 – Fluxo do modelo" in texts, str(texts))
    except Exception as e:
        report("set_figure_caption", False, str(e))
        traceback.print_exc()


# ────────────────────────────────────────────────────────
# Test 13: set_figure_source
# ────────────────────────────────────────────────────────
def test_set_figure_source():
    path = os.path.join(TMPDIR, "test_figure_source.docx")
    out_path = os.path.join(TMPDIR, "test_figure_source_out.docx")
    make_docx_with_inline_figure(path)
    try:
        result = sd.set_figure_source(path, 1, "Autor (2024)", out_path)
        report("set_figure_source: ok", result.get("ok", False), str(result))
        doc2 = Document(out_path)
        texts = [sd._get_paragraph_full_text(p).strip() for p in doc2.paragraphs if sd._get_paragraph_full_text(p).strip()]
        report("set_figure_source: prefixes source", "Fonte: Autor (2024)" in texts, str(texts))
    except Exception as e:
        report("set_figure_source", False, str(e))
        traceback.print_exc()


# ────────────────────────────────────────────────────────
# Test 14: validate_tables_with_comments
# ────────────────────────────────────────────────────────
def test_validate_tables_with_comments():
    path = os.path.join(TMPDIR, "test_vtc.docx")
    out_path = os.path.join(TMPDIR, "test_vtc_out.docx")
    make_simple_docx(path)
    try:
        result = sd.validate_tables_with_comments(path, out_path)
        report("validate_tables_with_comments: returns dict", isinstance(result, dict))
        report("validate_tables_with_comments: has report", "report" in result)
    except Exception as e:
        report("validate_tables_with_comments", False, str(e))
        traceback.print_exc()


# ────────────────────────────────────────────────────────
# Test 15: edge case — empty doc
# ────────────────────────────────────────────────────────
def test_edge_empty_doc():
    path = os.path.join(TMPDIR, "test_empty.docx")
    doc = Document()
    doc.save(path)
    try:
        result = sd.renumber_tables(path)
        report("edge: renumber_tables on empty doc", result.get("ok") is False, str(result.get("message", "")))

        result2 = sd.list_comments(path)
        report("edge: list_comments on empty doc", isinstance(result2, list), str(result2))

        result3 = sd.renumber_equations(path)
        report("edge: renumber_equations on empty doc", result3.get("ok") is False)
    except Exception as e:
        report("edge: empty doc", False, str(e))
        traceback.print_exc()


# ────────────────────────────────────────────────────────
# Test 16: paragraph_index out of range
# ────────────────────────────────────────────────────────
def test_edge_out_of_range():
    path = os.path.join(TMPDIR, "test_oor.docx")
    doc = Document()
    doc.add_paragraph("Only one paragraph.")
    doc.save(path)
    try:
        result = sd.replace_paragraph_text(path, 999, "Nope", "test")
        report("edge: out of range replace", "fora do range" in result.lower(), result)

        result2 = sd.insert_citation(path, 999, "(X)", "ref")
        report("edge: out of range citation", "fora do range" in result2.lower(), result2)

        result3 = sd.insert_candidate_comment(path, 999, "test")
        report("edge: out of range comment", "fora do range" in result3.lower(), result3)
    except Exception as e:
        report("edge: out of range", False, str(e))
        traceback.print_exc()


# ────────────────────────────────────────────────────────
# Run all tests
# ────────────────────────────────────────────────────────
if __name__ == "__main__":
    print(f"Temp dir: {TMPDIR}\n")

    test_list_comments_context()
    test_removed_public_apis()
    test_list_comments_resolved_by_thumbsup()
    test_reply_comment()
    test_replace_paragraph_text()
    test_remove_resolved_comments()
    test_remove_resolved_comments_with_resolved()
    test_remove_resolved_comments_with_thumbsup()
    test_renumber_tables()
    test_insert_citation()
    test_output_path_optional()
    test_renumber_equations()
    test_renumber_equations_with_orphan_reference()
    test_renumber_equations_with_duplicate_numbers()
    test_apply_table_style()
    test_set_table_caption()
    test_set_table_source()
    test_set_figure_caption()
    test_set_figure_source()
    test_validate_tables_with_comments()
    test_edge_empty_doc()
    test_edge_out_of_range()

    print(f"\n{'='*60}")
    passed = sum(1 for _, ok in RESULTS if ok)
    total = len(RESULTS)
    failed = total - passed
    print(f"  TOTAL: {total}  |  PASSED: {passed}  |  FAILED: {failed}")
    if failed:
        print("\n  Failed tests:")
        for name, ok in RESULTS:
            if not ok:
                print(f"    - {name}")
    print(f"{'='*60}")

    # Cleanup
    shutil.rmtree(TMPDIR, ignore_errors=True)

    sys.exit(0 if failed == 0 else 1)
